#!/usr/bin/env python3
"""Smoke test for the Python `python-docx` package (import name: docx).

Fully self-contained: no input files, no network, no packages beyond python-docx
(and its implied deps). Exercises the core API surface and exits 0 only if every
check passes, so it can be used as a pass/fail library validator:

    python3 python-docx.py

Install: pip install python-docx   (import name: docx)

This is the Python counterpart of data.table.R — same contract: a hard
not-installed guard (exit 1), a per-test harness that isolates failures, and a
PASS/exit-0 vs FAIL/exit-1 summary.
"""
import os
import sys

# This file is named after the package it tests, so it sits next to (and would
# shadow) the real top-level module. Drop this script's own directory from the
# import path before importing the package under test.
_here = os.path.dirname(os.path.abspath(__file__))
sys.path = [p for p in sys.path if p not in ("", ".") and os.path.abspath(p) != _here]

try:
    import docx
except ImportError:
    print("FAIL: package 'python-docx' is not installed")
    sys.exit(1)

import tempfile

from docx import Document


def _version(mod, dist):
    """Best-effort version string: module.__version__, else installed metadata."""
    v = getattr(mod, "__version__", None)
    if v:
        return v
    try:
        import importlib.metadata as m

        return m.version(dist)
    except Exception:
        return "unknown"


print(f"python-docx version: {_version(docx, 'python-docx')}")

failures = 0


def run_test(name, fn):
    """Run one check; a raised exception is a failure, not a crash."""
    global failures
    try:
        fn()
    except Exception as e:  # noqa: BLE001 - any failure is a test failure
        failures += 1
        print(f"  FAIL {name}: {e}")
    else:
        print(f"  ok   {name}")


def test_add_heading_and_paragraph():
    doc = Document()
    doc.add_heading("Title", 0)
    doc.add_paragraph("body")
    # add_heading(level=0) and add_paragraph each append one paragraph.
    assert len(doc.paragraphs) == 2
    assert doc.paragraphs[0].text == "Title"
    assert doc.paragraphs[1].text == "body"


def test_paragraph_run_text():
    doc = Document()
    p = doc.add_paragraph("hello ")
    p.add_run("world")
    assert p.text == "hello world"
    assert len(p.runs) == 2


def test_save_reopen_roundtrip_tempfile():
    doc = Document()
    doc.add_heading("Persisted Title", 0)
    doc.add_paragraph("first")
    doc.add_paragraph("second")
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        doc.save(path)
        assert os.path.getsize(path) > 0
        reopened = Document(path)
        assert len(reopened.paragraphs) == 3
        assert reopened.paragraphs[0].text == "Persisted Title"
        assert reopened.paragraphs[1].text == "first"
        assert reopened.paragraphs[2].text == "second"
    finally:
        os.remove(path)


run_test("add heading + paragraph", test_add_heading_and_paragraph)
run_test("paragraph run text", test_paragraph_run_text)
run_test("save/reopen roundtrip tempfile", test_save_reopen_roundtrip_tempfile)

if failures > 0:
    print(f"FAIL: {failures} test(s) failed")
    sys.exit(1)
print("PASS: all python-docx smoke tests passed")
